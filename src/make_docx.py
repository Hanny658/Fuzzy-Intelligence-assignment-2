"""Build report/report.docx from report/report.tex with pandoc (pypandoc_binary), content-identical.

    python src/make_docx.py

pandoc handles the LaTeX -> OMML maths, tables, figures and headings; the things it cannot number by
itself in docx (\\ref to tables/figures/sections, float captions, \\cite with a hand-written
thebibliography) are resolved here to the same numbers LaTeX produces, so the two documents read the same.
"""
from __future__ import annotations

import os
import re
import sys

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
REPORT = os.path.join(ROOT, "report")
TEX = os.path.join(REPORT, "report.tex")
DOCX = os.path.join(REPORT, "report.docx")

BS = "\\"  # a single backslash, to keep the replacement strings below readable


def number_labels(src: str) -> dict:
    """Map every \\label to the number LaTeX assigns it (tables, figures, sections in document order)."""
    numbers, counters = {}, {"table": 0, "figure": 0}
    sec = [0, 0]
    token = re.compile(r"\\section\{|\\subsection\{|\\begin\{table\}|\\begin\{figure\}|\\label\{([^}]+)\}")
    current = None
    for m in token.finditer(src):
        t = m.group(0)
        if t == BS + "section{":
            sec = [sec[0] + 1, 0]
            current = f"{sec[0]}"
        elif t == BS + "subsection{":
            sec[1] += 1
            current = f"{sec[0]}.{sec[1]}"
        elif t == BS + "begin{table}":
            counters["table"] += 1
            current = str(counters["table"])
        elif t == BS + "begin{figure}":
            counters["figure"] += 1
            current = str(counters["figure"])
        else:
            numbers[m.group(1)] = current
    return numbers


def table_text(s: str) -> str:
    """Turn the maths inside table cells into plain text so Word gets ordinary numbers."""
    s = s.replace("$" + BS + "pm$", "±").replace(BS + "pm", " ± ")
    s = re.sub(r"\$(-?\d+\.\d+)\$", lambda m: m.group(1).replace("-", "\u2212"), s)
    s = re.sub(r"\$(\d\.\d+) ± (\d\.\d+)\$", r"\1 ± \2", s)
    s = re.sub(r"^\\cmidrule.*$", "", s, flags=re.M)
    return s


def preprocess(src: str) -> str:
    labels = number_labels(src)
    src = re.sub(r"~?\\ref\{([^}]+)\}", lambda m: " " + labels[m.group(1)], src)

    # bibliography: \bibitem order -> [n]
    keys = re.findall(r"\\bibitem\{([^}]+)\}", src)
    order = {k: i + 1 for i, k in enumerate(keys)}
    src = re.sub(r"~?\\cite\{([^}]+)\}",
                 lambda m: " [" + ", ".join(str(order[k.strip()]) for k in m.group(1).split(",")) + "]", src)
    src = src.replace("  ", " ")
    body = re.search(r"\\begin\{thebibliography\}\{\d+\}(.*?)\\end\{thebibliography\}", src, re.S)
    if body:
        items = re.split(r"\\bibitem\{[^}]+\}", body.group(1))[1:]
        # a leading "{}" keeps pandoc from reading "[n]" as an optional argument of the previous command
        refs = BS + "section*{References}\n" + "\n".join(f"{{}}[{i + 1}]~{it.strip()}\n" for i, it in enumerate(items))
        src = src[:body.start()] + refs + src[body.end():]

    # pandoc does not number floats in docx: prefix captions with the numbers LaTeX assigns
    counters = {"table": 0, "figure": 0}

    def number_caption(m):
        kind = m.group(1)
        counters[kind] += 1
        prefix = f"{kind.capitalize()} {counters[kind]}: "
        return m.group(0).replace(BS + "caption{", BS + "caption{" + prefix, 1)

    src = re.sub(r"\\begin\{(table|figure)\}.*?\\end\{\1\}", number_caption, src, flags=re.S)

    # run-in paragraph headings -> bold text (pandoc would otherwise make numbered level-4 headings)
    src = re.sub(r"\\paragraph\{([^}]*)\}", BS + BS + "textbf{" + r"\1" + "} ", src)

    # generated tables: inline the files and convert cell maths to text
    def inline_input(m):
        return table_text(open(os.path.join(REPORT, m.group(1)), encoding="utf-8").read())

    src = re.sub(r"\\input\{([^}]+)\}", inline_input, src)
    # hand-written paradigm table: collapse the two-row multicolumn header into one row
    src = re.sub(
        r" & \\multicolumn\{2\}\{c\}\{NUH-g2\}.*?Training rule & CV AUC & CV EER & CV AUC & CV EER & NUH & WDBC \\\\",
        lambda m: "Training rule & NUH CV AUC & NUH CV EER & WDBC CV AUC & WDBC CV EER & Fit NUH (s) & Fit WDBC (s) " + BS + BS,
        src, flags=re.S)
    src = re.sub(r"\\begin\{tabular\}\{lcccccc\}.*?\\end\{tabular\}", lambda m: table_text(m.group(0)), src, flags=re.S)
    src = src.replace("$5{" + BS + "times}5$", "5×5")

    # layout-only commands pandoc does not need
    src = re.sub(r"\\setlength\{\\tabcolsep\}\{[^}]*\}", "", src)
    src = src.replace(BS + "centering" + BS + "footnotesize", "").replace(BS + "footnotesize", "")
    src = src.replace(BS + "hfill", " ")
    return src


def format_docx(path: str) -> None:
    """Apply the report's A4 layout and a few Word-specific pagination safeguards."""
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Mm, Pt

    doc = Document(path)
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(22)
        section.bottom_margin = Mm(22)
        section.left_margin = Mm(22)
        section.right_margin = Mm(22)
        section.footer_distance = Mm(10)

        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instruction = OxmlElement("w:instrText")
        instruction.set(qn("xml:space"), "preserve")
        instruction.text = " PAGE "
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.extend((begin, instruction, end))

    widths_mm = {
        0: (42, 26, 28, 35, 35),
        1: (38, 32, 40, 28, 28),
        2: (42, 26, 28, 35, 35),
    }
    for table_index, table in enumerate(doc.tables):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        if table.rows:
            table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
        column_widths = widths_mm.get(table_index)
        for row in table.rows:
            row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
            for column_index, cell in enumerate(row.cells):
                if column_widths and column_index < len(column_widths):
                    cell.width = Mm(column_widths[column_index])
                    tc_width = cell._tc.get_or_add_tcPr().get_or_add_tcW()
                    tc_width.type = "dxa"
                    tc_width.w = int(Mm(column_widths[column_index]).twips)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith(("Table ", "Figure ")):
            paragraph.paragraph_format.keep_with_next = True
        if text.startswith("Table 2:"):
            paragraph.paragraph_format.page_break_before = True
        if text == "References":
            paragraph.paragraph_format.page_break_before = True

    doc.save(path)


def main() -> None:
    import pypandoc

    src = open(TEX, encoding="utf-8").read()
    tmp = os.path.join(REPORT, "_report_docx.tex")
    open(tmp, "w", encoding="utf-8").write(preprocess(src))
    cwd = os.getcwd()
    os.chdir(REPORT)
    try:
        pypandoc.convert_file(
            "_report_docx.tex", "docx", outputfile="report.docx",
            extra_args=["--resource-path=.;../figures;../results", "--number-sections", "--dpi=200"],
        )
        format_docx(DOCX)
    finally:
        os.chdir(cwd)
        os.remove(tmp)
    print("wrote", DOCX, f"({os.path.getsize(DOCX) // 1024} KB)")


if __name__ == "__main__":
    sys.exit(main())
